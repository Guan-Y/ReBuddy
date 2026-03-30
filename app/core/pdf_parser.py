"""
PDF解析核心模块 - 从原paper_parser.py迁移并重构
去除Flask依赖，保持纯算法逻辑
支持 PyMuPDF 和 unstructured 两种解析引擎
"""

import os
import fitz  # PyMuPDF
from typing import List, Optional, Dict, Any
import json
import pathlib
import csv
import re

from app.core.llm_client import LLMClient
from app.config import Config
from app.models.schemas import PaperMetadata

from app.core.prompts.papers_extract import PAPER_EXTRACT_PROMPT

# Unstructured 导入（用于智能解析）
try:
    from unstructured.partition.pdf import partition_pdf
    UNSTRUCTURED_AVAILABLE = True
    print("✅ Unstructured 库已安装，将使用智能解析模式")

    import os
# 强制指定 OCR 引擎为 paddle
    os.environ["OCR_AGENT"] = "unstructured.partition.utils.ocr_models.paddle_ocr.OCRAgentPaddle"
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    print("⚠️ Unstructured 库未安装，将使用 PyMuPDF 传统模式")

# pypdfplumber 导入（用于快速解析）
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    print("✅ pypdfplumber 库已安装，将支持快速解析模式")
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("⚠️ pypdfplumber 库未安装，快速解析模式不可用")

# PaddleOCR 导入（用于 PaddleOCR + unstructured 组合解析方案）
try:
    from paddleocr import PaddleOCR
    from paddleocr import PPStructureV3 as PPStructure
    print("Using PPStructureV3")
    PADDLEOCR_AVAILABLE = True
    print("✅ PaddleOCR 已安装，支持 PaddleOCR + unstructured 组合解析方案")
except ImportError:
    PADDLEOCR_AVAILABLE = False
    print("⚠️ PaddleOCR 未安装，PaddleOCR + unstructured 组合解析方案不可用")

# Unstructured 元素类导入（用于创建标准元素对象）
try:
    from unstructured.documents.elements import (
        NarrativeText,
        Title,
        Header,
        ListItem,
        Table,
        Image,
        ElementMetadata
    )
    UNSTRUCTURED_ELEMENTS_AVAILABLE = True
    print("✅ Unstructured 元素类已安装，支持标准元素格式")
except ImportError:
    UNSTRUCTURED_ELEMENTS_AVAILABLE = False
    print("⚠️ Unstructured 元素类未安装，将使用简化元素格式")



# ==========================================
# 2. PDF 解析模块
# ==========================================
def table_to_markdown(table_data: List[List[str]]) -> str:
    """将表格数据转换为 Markdown 格式字符串，便于 LLM 理解"""
    if not table_data:
        return ""
    
    md_str = "\n\n| " + " | ".join(map(str, table_data[0])) + " |\n"
    md_str += "| " + " | ".join(["---"] * len(table_data[0])) + " |\n"
    
    for row in table_data[1:]:
        md_str += "| " + " | ".join(map(str, row)) + " |\n"
    
    return md_str + "\n"


def _extract_text_from_file(file_path: str, output_dir: str, file_ext: str) -> Dict:
    """
    从文本类文件中直接提取文本内容
    支持：md, docx, csv, xlsx, txt
    
    Args:
        file_path: 文件路径
        output_dir: 输出目录
        file_ext: 文件扩展名
        
    Returns:
        {
            "text": 完整文本,
            "images": [],
            "tables": []
        }
    """
    print(f"📄 检测到文本类文件 ({file_ext})，直接提取文本内容")
    
    full_text = ""
    
    try:
        if file_ext == '.md':
            # Markdown 文件
            with open(file_path, 'r', encoding='utf-8') as f:
                full_text = f.read()
            print(f"✅ Markdown 文件读取完成: {len(full_text)} 字符")
            
        elif file_ext == '.txt':
            # 纯文本文件
            with open(file_path, 'r', encoding='utf-8') as f:
                full_text = f.read()
            print(f"✅ 纯文本文件读取完成: {len(full_text)} 字符")
            
        elif file_ext == '.csv':
            # CSV 文件
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
                # 将 CSV 转换为 Markdown 表格格式
                if rows:
                    full_text = table_to_markdown(rows)
            print(f"✅ CSV 文件读取完成: {len(rows)} 行")
            
        elif file_ext == '.xlsx':
            # Excel 文件
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True)
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    full_text += f"\n## Sheet: {sheet_name}\n\n"
                    # 将 Excel 表格转换为 Markdown
                    table_data = []
                    for row in sheet.iter_rows(values_only=True):
                        table_data.append([str(cell) if cell is not None else "" for cell in row])
                    if table_data:
                        full_text += table_to_markdown(table_data)
                wb.close()
                print(f"✅ Excel 文件读取完成: {len(wb.sheetnames)} 个工作表")
            except ImportError:
                print("⚠️ openpyxl 未安装，尝试使用 xlrd")
                try:
                    import xlrd
                    workbook = xlrd.open_workbook(file_path)
                    for sheet_name in workbook.sheet_names():
                        sheet = workbook.sheet_by_name(sheet_name)
                        full_text += f"\n## Sheet: {sheet_name}\n\n"
                        table_data = []
                        for row_idx in range(sheet.nrows):
                            row = []
                            for col_idx in range(sheet.ncols):
                                cell_value = sheet.cell_value(row_idx, col_idx)
                                row.append(str(cell_value))
                            table_data.append(row)
                        if table_data:
                            full_text += table_to_markdown(table_data)
                    print(f"✅ Excel 文件读取完成: {len(workbook.sheet_names())} 个工作表")
                except ImportError:
                    print("❌ xlrd 也未安装，无法读取 Excel 文件")
                    full_text = "Error: 无法读取 Excel 文件，请安装 openpyxl 或 xlrd"
                    
        elif file_ext == '.docx':
            # Word 文档
            try:
                from docx import Document
                doc = Document(file_path)
                for para in doc.paragraphs:
                    full_text += para.text + "\n"
                # 提取表格
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        table_data.append([cell.text for cell in row.cells])
                    if table_data:
                        full_text += table_to_markdown(table_data)
                print(f"✅ Word 文档读取完成: {len(doc.paragraphs)} 个段落, {len(doc.tables)} 个表格")
            except ImportError:
                print("❌ python-docx 未安装，无法读取 Word 文档")
                full_text = "Error: 无法读取 Word 文档，请安装 python-docx"
        
        return {
            "text": full_text,
            "images": [],
            "tables": []
        }
        
    except Exception as e:
        print(f"❌ 文本文件读取失败: {e}")
        return {
            "text": f"Error: 读取文件失败 - {str(e)}",
            "images": [],
            "tables": []
        }


def save_image(page, xref, output_dir, page_num, img_idx):
    """辅助函数：保存提取的图片"""
    try:
        pix = fitz.Pixmap(page.parent, xref)
        # 如果是 CMYK 等格式，转换为 RGB
        if pix.n - pix.alpha > 3:
            pix = fitz.Pixmap(fitz.csRGB, pix)
        
        filename = f"page_{page_num+1}_img_{img_idx}.png"
        filepath = os.path.join(output_dir, filename)
        pix.save(filepath)
        pix = None
        return filepath
    except Exception as e:
        print(f"⚠️ 保存图片失败 (Page {page_num+1}): {e}")
        return None


def process_pdf(pdf_path: str, output_dir: str = Config.PARSED_PAPERS_PATH) -> Dict:
    """
    处理 PDF：提取文本、保存图片、保存表格
    对于 md、docx、csv、xlsx 和 txt 文件直接提取文本内容
    返回字典包含：full_text, image_paths, table_paths
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"文件未找到: {pdf_path}")
    
    # 获取文件扩展名
    file_ext = pathlib.Path(pdf_path).suffix.lower()
    
    # 支持直接文本提取的文件类型
    text_based_extensions = {'.md', '.docx', '.csv', '.xlsx', '.txt'}
    
    # 如果是文本类文件，直接提取文本
    if file_ext in text_based_extensions:
        return _extract_text_from_file(pdf_path, output_dir, file_ext)
    
    # 否则使用 PyMuPDF 解析 PDF
    doc = fitz.open(pdf_path)
    
    # 准备输出目录
    
    pdf_name = pathlib.Path(pdf_path).stem
    assets_dir = os.path.join(output_dir, pdf_name)
    images_dir = os.path.join(assets_dir, "images")
    tables_dir = os.path.join(assets_dir, "tables")
    
    os.makedirs(images_dir, exist_ok=True)
    os.makedirs(tables_dir, exist_ok=True)
    
    full_text_with_context = ""
    extracted_images = []
    extracted_tables = []
    
    print(f"📂 正在解析 PDF 内容及资产，输出目录: {assets_dir}")

    for page_num, page in enumerate(doc):
        # 1. 提取基础文本
        page_text = page.get_text()
        
        # 2. 提取并保存图片
        image_list = page.get_images(full=True)
        for img_idx, img in enumerate(image_list):
            xref = img[0]
            img_path = save_image(page, xref, images_dir, page_num, img_idx)
            if img_path:
                extracted_images.append(img_path)
                # 在文本中插入标记，告诉 LLM 这里有张图
                page_text += f"\n[Image Extracted: {os.path.basename(img_path)}]\n"

        # 3. 提取表格
        try:
            tables = page.find_tables()
            if tables:
                for tab_idx, table in enumerate(tables):
                    table_data = table.extract()
                    
                    # A. 保存为 CSV
                    csv_filename = f"page_{page_num+1}_tab_{tab_idx}.csv"
                    csv_filepath = os.path.join(tables_dir, csv_filename)
                    with open(csv_filepath, 'w', newline='', encoding='utf-8') as f:
                        writer = csv.writer(f)
                        writer.writerows(table_data)
                    extracted_tables.append(csv_filepath)
                    
                    # B. 转换为 Markdown 插入文本流
                    md_table = table_to_markdown(table_data)
                    page_text += f"\n[Table Detected and Parsed]:\n{md_table}\n"
        except AttributeError:
            print("⚠️ 当前 PyMuPDF 版本不支持 find_tables，跳过表格提取。")
        except Exception as e:
            print(f"⚠️ 表格提取出错 (Page {page_num+1}): {e}")

        full_text_with_context += f"--- Page {page_num + 1} ---\n{page_text}\n"
        
    doc.close()
    
    return {
        "text": full_text_with_context,
        "images": extracted_images,
        "tables": extracted_tables
    }


# ==========================================
# 2.5. Unstructured 智能解析模块
# ==========================================
def process_pdf_with_unstructured(
    pdf_path: str,
    output_dir: str = Config.PARSED_PAPERS_PATH,
    use_ocr: bool = False,
    ocr_engine: str = "none",
    use_fast: bool = True,
    use_paddle_unstructured: bool = False,
    use_ppstructure: bool = False
) -> Dict:
    """
    使用 unstructured 库智能解析 PDF 或使用 pypdfplumber 快速解析
    支持智能分段、表格识别、图片提取

    Args:
        pdf_path: PDF 文件路径
        output_dir: 输出目录
        use_ocr: 是否使用 OCR（预留）
        ocr_engine: OCR 引擎类型（"none", "paddleocr", "tesseract"）
        use_fast: 是否使用 pypdfplumber 快速解析模式（默认 False）
        use_paddle_unstructured: 是否使用 PaddleOCR + unstructured 元素类混合方案（推荐，默认 False）
        use_ppstructure: PaddleOCR 模式下是否使用 PP-Structure 进行表格识别（默认 True）

    Returns:
        {
            "text": 完整文本,
            "elements": 结构化元素列表,
            "structured_info": 结构化信息（标题、页码等）,
            "images": 图片路径列表,
            "tables": 表格元素列表
        }
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"文件未找到: {pdf_path}")

    # ✅ 优先级1：使用 PaddleOCR + unstructured 元素类混合方案（推荐）
    if use_paddle_unstructured:
        if not PADDLEOCR_AVAILABLE:
            print("⚠️ PaddleOCR 未安装，回退到 unstructured 模式")
        elif not UNSTRUCTURED_ELEMENTS_AVAILABLE:
            print("⚠️ Unstructured 元素类未安装，回退到 unstructured 模式")
        else:
            return _process_pdf_with_paddleocr_unstructured(pdf_path, output_dir)

    # ✅ 优先级2：使用 pypdfplumber 快速解析模式
    if use_fast:
        if not PDFPLUMBER_AVAILABLE:
            print("⚠️ pypdfplumber 未安装，回退到 unstructured 模式")
        else:
            return _process_pdf_with_pdfplumber(pdf_path, output_dir)

    if not UNSTRUCTURED_AVAILABLE:
        print("⚠️ Unstructured 未安装，回退到 PyMuPDF 模式")
        return process_pdf(pdf_path, output_dir)
    
    # 准备输出目录
    pdf_name = pathlib.Path(pdf_path).stem
    assets_dir = os.path.join(output_dir, pdf_name)
    images_dir = os.path.join(assets_dir, "images")
    
    os.makedirs(images_dir, exist_ok=True)
    
    print(f"📂 正在使用 Unstructured 解析 PDF: {pdf_path}")
    
    try:
        # 使用 unstructured 解析 PDF
        elements = partition_pdf(
            filename=pdf_path,
            # strategy="fast",  # 默认fast；hi_res高分辨率策略（更准确）
            extract_tables=True,  # 提取表格
            extract_images_in_pdf=True,  # 提取图片
            extract_image_block_to_payload=True,  # 提取图片块
            infer_table_structure=True,  # 推断表格结构
            chunking_strategy="by_title",  # ✅ 智能分段（按标题）
            max_characters=1024,  # 最大字符数（与 kb_manager 一致）
            combine_under_n_chars=200,  # 小于 200 字符的块合并
            new_after_n_chars=1500,  # 超过 1500 字符强制分段
            # OCR 配置（预留）
            ocr_languages=["eng", "chi_sim"] if use_ocr else None,
            extract_image_block_output_dir=images_dir if use_ocr else None,
        )
        
        print(f"✅ Unstructured 解析完成，提取 {len(elements)} 个元素")
        
        # 1. 构建完整文本（用于 LLM 分析）
        full_text = ""
        for element in elements:
            full_text += element.text + "\n"
        
        # 2. 提取结构化信息（增强元数据）
        structured_info = {
            "titles": [el.text for el in elements if el.category == "Title"],
            "headers": [el.text for el in elements if el.category == "Header"],
            "narrative_texts": [el.text for el in elements if el.category == "NarrativeText"],
            "list_items": [el.text for el in elements if el.category == "ListItem"],
            "page_numbers": list(set([el.metadata.page_number for el in elements if hasattr(el.metadata, 'page_number')])),
            "table_count": len([el for el in elements if el.category == "Table"]),
            "image_count": len([el for el in elements if el.category == "Image"]),
            "total_elements": len(elements),
        }
        
        # 3. 提取图片路径
        images = []
        for element in elements:
            if element.category == "Image":
                if hasattr(element.metadata, 'image_path') and element.metadata.image_path:
                    images.append(element.metadata.image_path)
        
        # 4. 提取表格元素
        tables = [el for el in elements if el.category == "Table"]
        
        print(f"   📊 统计: {len(images)} 张图片, {len(tables)} 个表格")
        print(f"   📝 结构: {len(structured_info['titles'])} 个标题, {len(structured_info['narrative_texts'])} 段正文")
        
        return {
            "text": full_text,
            "elements": elements,
            "structured_info": structured_info,
            "images": images,
            "tables": tables
        }
        
    except Exception as e:
        print(f"❌ Unstructured 解析失败: {e}")
        print("🔄 回退到 PyMuPDF 模式")
        return process_pdf(pdf_path, output_dir)


def extract_element_metadata(element) -> Dict[str, Any]:
    """
    提取单个元素的元数据
    
    Args:
        element: unstructured 元素对象
        
    Returns:
        元数据字典
    """
    metadata = {
        "category": element.category,
        "text": element.text[:200] if len(element.text) > 200 else element.text,
    }
    
    # 提取页码
    if hasattr(element.metadata, 'page_number'):
        metadata["page_number"] = element.metadata.page_number
    
    # 提取父元素 ID
    if hasattr(element.metadata, 'parent_id'):
        metadata["parent_id"] = element.metadata.parent_id
    
    # 提取坐标
    if hasattr(element.metadata, 'coordinates'):
        coords = element.metadata.coordinates
        if coords:
            metadata["coordinates"] = {
                "points": coords.get("points", []),
                "system": coords.get("system", ""),
            }
    
    # 图片特有元数据
    if element.category == "Image":
        if hasattr(element.metadata, 'image_path'):
            metadata["image_path"] = element.metadata.image_path
        if hasattr(element.metadata, 'caption'):
            metadata["caption"] = element.metadata.caption
    
    # 表格特有元数据
    if element.category == "Table":
        if hasattr(element.metadata, 'text_as_html'):
            metadata["table_html"] = element.metadata.text_as_html[:500]  # 截断
        if hasattr(element.metadata, 'table_cells'):
            metadata["table_cells_count"] = len(element.metadata.table_cells)
    
    return metadata


# ==========================================
# 3. LLM 提取模块
# ==========================================
def extract_json_blocks(text):
    """从文本中提取 ```json ... ``` 格式的JSON代码块"""
    pattern = r'```json\s*([\s\S]*?)\s*```'
    matches = re.findall(pattern, text)
    
    results = []
    for match in matches:
        try:
            json_obj = json.loads(match)
            results.append(match)
        except json.JSONDecodeError as e:
            print(f"JSON解析失败: {e}")
            continue
    
    return results


async def analyze_paper_with_llm(pdf_content: Dict, llm_client: LLMClient = None, max_retries: int = 2) -> PaperMetadata:
    """
    调用模型进行深度解析
    支持使用结构化信息增强提取准确性
    """
    if llm_client is None:
        llm_client = LLMClient()

    # ✅ 提取结构化信息（如果存在）
    structured_info = pdf_content.get('structured_info', {})
    structured_info_section = ""
    
    if structured_info:
        titles = structured_info.get('titles', [])
        headers = structured_info.get('headers', [])
        table_count = structured_info.get('table_count', 0)
        image_count = structured_info.get('image_count', 0)
        
        structured_info_section = f"""
【结构化信息】
- 论文标题: {titles[0] if titles else '未知'}
- 章节标题: {', '.join(headers[:5]) if headers else '无'}{'...' if len(headers) > 5 else ''}
- 表格数量: {table_count}
- 图片数量: {image_count}
- 总元素数: {structured_info.get('total_elements', 0)}
"""

    base_prompt = PAPER_EXTRACT_PROMPT.format(content=pdf_content['text'], structured_info=structured_info_section)

    print("🤖 正在请求 AI 专家进行分析...")

    for attempt in range(max_retries):
        try:
            if attempt > 0:
                prompt = base_prompt + f"\n\n【重试提示】这是第 {attempt + 1} 次尝试，请特别注意JSON格式的正确性。"
            else:
                prompt = base_prompt

            # 异步调用 LLM
            response = await llm_client.generate_async(prompt)
            
            json_content = extract_json_blocks(response.content)

            if not json_content:
                raise ValueError("未能提取到有效的JSON内容")

            metadata = PaperMetadata.model_validate_json(json_content[0])

            # 强制更新资产字段
            metadata.image_paths = pdf_content['images']
            metadata.table_paths = pdf_content['tables']
            metadata.asset_images_count = len(pdf_content['images'])
            metadata.asset_tables_count = len(pdf_content['tables'])

            # 自动设置 has_code
            metadata.has_code = bool(metadata.code_url)

            # 确保 summary 和 abstract 一致
            if not metadata.summary and metadata.abstract:
                metadata.summary = metadata.abstract
            elif not metadata.abstract and metadata.summary:
                metadata.abstract = metadata.summary

            print(f"✅ 解析成功 (尝试 {attempt + 1}/{max_retries})")
            return metadata

        except Exception as e:
            print(f"❌ 第 {attempt + 1} 次解析失败: {e}")
            if attempt < max_retries - 1:
                import asyncio
                await asyncio.sleep(3)
            else:
                print(f"💥 所有 {max_retries} 次尝试均失败")
                print("🛡️ 返回最小可用元数据对象")
                return PaperMetadata(
                    title="解析失败",
                    authors="",
                    year=0,
                    abstract="由于解析失败，无法生成摘要。"
                )

    return None


def analyze_paper_with_llm_sync(pdf_content: Dict, llm_client: LLMClient = None, max_retries: int = 2) -> PaperMetadata:
    """
    同步调用模型进行深度解析（用于旧的同步接口）
    注意：此函数在事件循环中运行异步函数
    """
    import asyncio

    async def _analyze():
        return await analyze_paper_with_llm(pdf_content, llm_client, max_retries)

    try:
        # 尝试获取当前事件循环
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # 如果事件循环正在运行，创建一个新的事件循环
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, _analyze())
                return future.result()
        else:
            # 如果事件循环没有运行，直接运行
            return asyncio.run(_analyze())
    except RuntimeError:
        # 如果没有事件循环，创建一个新的
        return asyncio.run(_analyze())


def parse_pdf_to_metadata(pdf_path: str,
                          file_id: str = None,
                          output_dir: str = Config.PARSED_PAPERS_PATH,
                          user_id: str = None) -> Dict:
    """
    解析PDF并返回元数据
    Args:
        pdf_path: PDF文件路径
        file_id: 文件ID
        output_dir: 输出目录
        user_id: 用户ID
    Returns:
        解析结果字典
    """
    if output_dir is None:
        raise ValueError("output_dir 参数必须提供")

    if user_id is None:
        raise ValueError("user_id 参数必须提供")

    try:
        # 1. 解析 PDF
        content = process_pdf(pdf_path, output_dir)
        print(f"✅ 解析完成。文本长度: {len(content['text'])} 字符")
        print(f"📦 提取资源: {len(content['images'])} 张图片, {len(content['tables'])} 个表格")

        # 2. LLM 分析（使用同步版本）
        metadata = analyze_paper_with_llm_sync(content)

        if metadata:
            result = metadata.model_dump()

            # 保存元数据到文件
            pdf_name = pathlib.Path(pdf_path).stem
            metadata_path = os.path.join(output_dir, pdf_name, "metadata.json")

            os.makedirs(os.path.dirname(metadata_path), exist_ok=True)
            with open(metadata_path, 'w', encoding='utf-8') as f:
                f.write(json.dumps(result, indent=2, ensure_ascii=False))

            # 将元数据存入用户专属的知识库
            try:
                from app.core.kb_manager import get_kb_manager
                kb_manager = get_kb_manager(user_id=user_id)
                kb_manager.add_paper(full_text=content['text'], metadata_obj=metadata, file_id=file_id)
                print(f"✅ 论文已存入用户 {user_id} 的知识库")
            except Exception as e:
                print(f"⚠️ 存入知识库失败: {e}")
                # 知识库存储失败不影响解析结果

            return {"status": "success", "data": result}
        else:
            return {"status": "error", "message": "解析失败"}

    except Exception as e:
        return {"status": "error", "message": str(e)}


# ==========================================
# 4. 新增：分步解析接口
# ==========================================
def process_pdf_for_knowledge_base(
    pdf_path: str,
    file_id: str,
    kb_id: str,
    user_id: str,
    output_dir: str = None,
    use_unstructured: bool = False,
    use_fast: bool = False,
    use_paddle_unstructured: bool = False,
    use_ppstructure: bool = False
) -> Dict:
    """
    为知识库处理 PDF 文件（分步接口）
    1. 提取文本（同步）
    2. 添加原文到向量库（同步）
    3. LLM 分析元数据（异步可选）
    4. 添加摘要到向量库（异步可选）

    Args:
        pdf_path: PDF 文件路径
        file_id: 文件ID
        kb_id: 知识库ID
        user_id: 用户ID
        output_dir: 输出目录
        use_unstructured: 是否使用 unstructured 智能解析（默认 True）
        use_fast: 是否使用 pypdfplumber 快速解析模式（默认 False）
        use_paddle_unstructured: 是否使用 PaddleOCR + unstructured 元素类混合方案（推荐，默认 False）
        use_ppstructure: PaddleOCR 模式下是否使用 PP-Structure 进行表格识别（默认 True）

    Returns:
        {'status': 'success', 'full_text': ..., 'chunk_count': ..., 'images': ..., 'tables': ..., 'elements': ..., 'structured_info': ...}
    """
    from app.core.kb_manager import get_kb_manager

    if output_dir is None:
        user_paths = Config.get_user_paths(user_id)
        output_dir = user_paths['parsed_papers_path']

    try:
        # 1. 提取 PDF 文本、图片、表格（同步）
        if use_paddle_unstructured:
            # ✅ 使用 PaddleOCR + unstructured 元素类混合方案（推荐）
            content = process_pdf_with_unstructured(
                pdf_path,
                output_dir,
                use_paddle_unstructured=True,
                use_ppstructure=use_ppstructure
            )
            print(f"✅ PaddleOCR + Unstructured 元素类混合提取完成: {len(content['text'])} 字符")
        elif use_fast:
            # ✅ 使用 pypdfplumber 快速解析
            content = process_pdf_with_unstructured(pdf_path, output_dir, use_fast=True)
            print(f"✅ pypdfplumber PDF 快速提取完成: {len(content['text'])} 字符")
        elif use_unstructured and UNSTRUCTURED_AVAILABLE:
            # ✅ 使用 unstructured 智能解析
            content = process_pdf_with_unstructured(pdf_path, output_dir)
            print(f"✅ Unstructured PDF 提取完成: {len(content['text'])} 字符")
        else:
            # 回退到 PyMuPDF 传统模式
            content = process_pdf(pdf_path, output_dir)
            print(f"✅ PyMuPDF PDF 提取完成: {len(content['text'])} 字符")

        # 2. 添加原文到向量库（同步，原文切分后可直接使用）
        kb_manager = get_kb_manager(user_id=user_id)
        
        # ✅ 传递 elements 参数以支持智能分段
        elements = content.get('elements', None)
        chunk_count = kb_manager._add_paper_fulltext_qdrant(
            full_text=content['text'],
            file_id=file_id,
            kb_id=kb_id,
            elements=elements  # ✅ 传递结构化元素
        )
        print(f"✅ 原文已添加到向量库: {chunk_count} 个片段")

        result = {
            'status': 'success',
            'full_text': content['text'],
            'chunk_count': chunk_count,
            'images': content['images'],
            'tables': content['tables']
        }
        
        # ✅ 添加结构化信息（如果存在）
        if 'elements' in content:
            result['elements'] = content['elements']
        if 'structured_info' in content:
            result['structured_info'] = content['structured_info']
        
        return result

    except Exception as e:
        print(f"❌ PDF 处理失败: {e}")
        return {'status': 'error', 'message': str(e)}


async def analyze_and_add_summary(
    pdf_content: Dict,
    file_id: str,
    kb_id: str,
    user_id: str
) -> Dict:
    """
    异步分析 PDF 并添加摘要到向量库
    Args:
        pdf_content: PDF 解析结果（包含 text, images, tables）
        file_id: 文件ID
        kb_id: 知识库ID
        user_id: 用户ID
    Returns:
        {'status': 'success', 'summary_id': ..., 'metadata': ...}
    """
    from app.core.kb_manager import get_kb_manager
    from app.core.llm_client import LLMClient

    try:
        # 1. 异步调用 LLM 进行元数据分析
        llm_client = LLMClient()
        metadata = await analyze_paper_with_llm(pdf_content, llm_client)

        # 2. 添加摘要到向量库
        kb_manager = get_kb_manager(user_id=user_id)
        summary_id = kb_manager._add_paper_summary_qdrant(
            metadata_obj=metadata,
            file_id=file_id,
            kb_id=kb_id
        )
        print(f"✅ 摘要已添加到向量库: {summary_id}")

        return {
            'status': 'success',
            'summary_id': summary_id,
            'metadata': metadata.model_dump()
        }

    except Exception as e:
        print(f"❌ 摘要处理失败: {e}")
        return {'status': 'error', 'message': str(e)}


# ==========================================
# 2.6. pypdfplumber 快速解析模块
# ==========================================
def _process_pdf_with_pdfplumber(pdf_path: str, output_dir: str = Config.PARSED_PAPERS_PATH) -> Dict:
    """
    使用 pypdfplumber 快速解析 PDF
    专注于文本和表格提取，速度快但功能相对简单
    
    Args:
        pdf_path: PDF 文件路径
        output_dir: 输出目录
        
    Returns:
        {
            "text": 完整文本,
            "elements": 结构化元素列表（简化版）,
            "structured_info": 结构化信息,
            "images": 图片路径列表（空，pypdfplumber 不支持图片提取）,
            "tables": 表格元素列表
        }
    """
    if not PDFPLUMBER_AVAILABLE:
        print("⚠️ pypdfplumber 未安装，回退到 PyMuPDF 模式")
        return process_pdf(pdf_path, output_dir)
    
    # 准备输出目录
    pdf_name = pathlib.Path(pdf_path).stem
    assets_dir = os.path.join(output_dir, pdf_name)
    tables_dir = os.path.join(assets_dir, "tables")
    
    os.makedirs(tables_dir, exist_ok=True)
    
    print(f"📂 正在使用 pypdfplumber 快速解析 PDF: {pdf_path}")
    
    full_text = ""
    extracted_tables = []
    elements = []
    titles = []
    headers = []
    narrative_texts = []
    list_items = []
    page_numbers = []
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_numbers.append(page_num + 1)
                
                # 提取文本
                page_text = page.extract_text() or ""
                full_text += f"--- Page {page_num + 1} ---\n{page_text}\n"
                
                # 创建简化的文本元素
                if page_text.strip():
                    # 简单的启发式规则：如果文本较短且全大写，可能是标题
                    lines = page_text.strip().split('\n')
                    for line in lines:
                        line = line.strip()
                        if len(line) > 0:
                            if len(line) < 100 and line.isupper():
                                titles.append(line)
                                elements.append(_create_simple_element(line, "Title", page_num + 1))
                            elif len(line) < 50:
                                headers.append(line)
                                elements.append(_create_simple_element(line, "Header", page_num + 1))
                            else:
                                narrative_texts.append(line)
                                elements.append(_create_simple_element(line, "NarrativeText", page_num + 1))
                
                # 提取表格
                tables = page.extract_tables()
                if tables:
                    for tab_idx, table in enumerate(tables):
                        if table:
                            # 保存为 CSV
                            csv_filename = f"page_{page_num+1}_tab_{tab_idx}.csv"
                            csv_filepath = os.path.join(tables_dir, csv_filename)
                            with open(csv_filepath, 'w', newline='', encoding='utf-8') as f:
                                writer = csv.writer(f)
                                writer.writerows(table)
                            extracted_tables.append(csv_filepath)
                            
                            # 转换为 Markdown
                            md_table = table_to_markdown(table)
                            full_text += f"\n[Table Detected and Parsed]:\n{md_table}\n"
                            
                            # 创建表格元素
                            table_text = "\n".join([" | ".join([str(cell) if cell else "" for cell in row]) for row in table])
                            elements.append(_create_simple_element(table_text, "Table", page_num + 1))
        
        # 构建结构化信息
        structured_info = {
            "titles": titles,
            "headers": headers,
            "narrative_texts": narrative_texts,
            "list_items": list_items,
            "page_numbers": page_numbers,
            "table_count": len(extracted_tables),
            "image_count": 0,  # pypdfplumber 不支持图片提取
            "total_elements": len(elements),
        }
        
        print(f"✅ pypdfplumber 解析完成，提取 {len(elements)} 个元素")
        print(f"   📊 统计: 0 张图片, {len(extracted_tables)} 个表格")
        print(f"   📝 结构: {len(titles)} 个标题, {len(narrative_texts)} 段正文")
        
        return {
            "text": full_text,
            "elements": elements,
            "structured_info": structured_info,
            "images": [],  # pypdfplumber 不支持图片提取
            "tables": extracted_tables
        }
        
    except Exception as e:
        print(f"❌ pypdfplumber 解析失败: {e}")
        print("🔄 回退到 PyMuPDF 模式")
        return process_pdf(pdf_path, output_dir)


def _create_simple_element(text: str, category: str, page_number: int):
    """
    创建简化的元素对象（模拟 unstructured 元素）
    
    Args:
        text: 元素文本
        category: 元素类型
        page_number: 页码
        
    Returns:
        简化的元素对象
    """
    class SimpleElement:
        def __init__(self, text, category, page_number):
            self.text = text
            self.category = category
            self.metadata = SimpleMetadata(page_number)
    
    class SimpleMetadata:
        def __init__(self, page_number):
            self.page_number = page_number
    
    return SimpleElement(text, category, page_number)


# ==========================================
# 2.7. PaddleOCR + Unstructured 元素类混合方案
# ==========================================
def _process_pdf_with_paddleocr_unstructured(
    pdf_path: str,
    output_dir: str = Config.PARSED_PAPERS_PATH,
    use_ppstructure: bool = False
) -> Dict:
    """
    使用 PaddleOCR 进行 OCR，但创建标准的 unstructured 元素对象
    这样既能享受 PaddleOCR 的轻量级优势，又能使用 unstructured 的元素分类能力
    
    Args:
        pdf_path: PDF 文件路径
        output_dir: 输出目录
        use_ppstructure: 是否使用 PP-Structure 进行表格识别（默认 True）
        
    Returns:
        {
            "text": 完整文本,
            "elements": unstructured 标准元素列表,
            "structured_info": 结构化信息,
            "images": 图片路径列表（空）,
            "tables": 表格元素列表
        }
    """
    if not PADDLEOCR_AVAILABLE:
        print("⚠️ PaddleOCR 未安装，回退到 PyMuPDF 模式")
        return process_pdf(pdf_path, output_dir)
    
    # 准备输出目录
    pdf_name = pathlib.Path(pdf_path).stem
    assets_dir = os.path.join(output_dir, pdf_name)
    images_dir = os.path.join(assets_dir, "images")
    tables_dir = os.path.join(assets_dir, "tables")
    
    os.makedirs(images_dir, exist_ok=True)
    os.makedirs(tables_dir, exist_ok=True)
    
    print(f"📂 正在使用 PaddleOCR + Unstructured 元素类混合解析: {pdf_path}")
    
    # 初始化 PaddleOCR（轻量级模型，无 YOLO 依赖）
    # import pdb; pdb.set_trace()
    ocr = PaddleOCR(
        use_angle_cls=True,
        lang="ch",
        # use_gpu=False,
        det_model_dir=None,
        rec_model_dir=None,
        cls_model_dir=None
    )
    
    # 初始化 PP-Structure（用于表格识别，使用 PicoDet 而非 YOLO）
    table_engine = None
    if use_ppstructure:
        table_engine = PPStructure()
        print(f"   ✅ PP-Structure 已启用，使用 PicoDet 进行版面分析")
    
    full_text = ""
    extracted_tables = []
    elements = []
    titles = []
    headers = []
    narrative_texts = []
    list_items = []
    page_numbers = []
    
    try:
        # 使用 PyMuPDF 将 PDF 转为图片
        doc = fitz.open(pdf_path)

        # import pdb; pdb.set_trace()
        
        for page_num, page in enumerate(doc):
            page_numbers.append(page_num + 1)
            
            # 将当前页转换为图片
            mat = fitz.Matrix(2, 2)  # 2倍缩放以提高 OCR 准确率
            pix = page.get_pixmap(matrix=mat)
            img_path = os.path.join(images_dir, f"page_{page_num + 1}.png")
            pix.save(img_path)
            pix = None
            
            # 方案1：使用 PP-Structure 进行完整的版面分析（包括表格）
            if use_ppstructure and table_engine:
                result = table_engine(img_path)

                # ✅ 正确：result 是列表，result[0] 是字典
                if result and len(result) > 0:
                    page_result = result[0]

                    # ✅ 从字典中提取识别结果
                    rec_texts = page_result.get('rec_texts', [])
                    rec_polys = page_result.get('rec_polys', [])
                    rec_scores = page_result.get('rec_scores', [])

                    # 遍历识别的文本行
                    for idx, text in enumerate(rec_texts):
                        if not text.strip():
                            continue

                        # 获取对应的文本框坐标
                        box = rec_polys[idx].tolist() if idx < len(rec_polys) else None

                        # 创建 ElementMetadata
                        metadata = ElementMetadata(
                            page_number=page_num + 1,
                            coordinates={"points": box, "system": "pixel"} if box else None
                        )

                        # 简单规则：根据字数和标点判断标题
                        if len(text) < 30 and "。" not in text and "，" not in text:
                            titles.append(text)
                            title_element = Title(text=text, metadata=metadata)
                            elements.append(title_element)
                        else:
                            narrative_texts.append(text)
                            text_element = NarrativeText(text=text, metadata=metadata)
                            elements.append(text_element)

                        full_text += f"{text}\n"
            
            else:
                # 方案2：使用纯 PaddleOCR（仅文字检测+识别）+ 简单规则
                result = ocr.ocr(img_path)

                # ✅ 正确：result 是列表，result[0] 是字典
                if result and len(result) > 0:
                    page_result = result[0]

                    # ✅ 从字典中提取识别结果
                    rec_texts = page_result.get('rec_texts', [])
                    rec_polys = page_result.get('rec_polys', [])
                    rec_scores = page_result.get('rec_scores', [])

                    # 遍历识别的文本行
                    for idx, text in enumerate(rec_texts):
                        if not text.strip():
                            continue

                        # 获取对应的文本框坐标
                        box = rec_polys[idx].tolist() if idx < len(rec_polys) else None

                        # 创建 ElementMetadata
                        metadata = ElementMetadata(
                            page_number=page_num + 1,
                            coordinates={"points": box, "system": "pixel"} if box else None
                        )

                        # 简单规则：根据字数和标点判断标题
                        if len(text) < 30 and "。" not in text and "，" not in text:
                            # 创建标准的 Title 元素
                            titles.append(text)
                            title_element = Title(text=text, metadata=metadata)
                            elements.append(title_element)
                        else:
                            # 创建标准的 NarrativeText 元素
                            narrative_texts.append(text)
                            text_element = NarrativeText(text=text, metadata=metadata)
                            elements.append(text_element)

                        full_text += f"{text}\n"

                full_text += f"\n--- Page {page_num + 1} ---\n"
        
        doc.close()
        
        # 构建结构化信息
        structured_info = {
            "titles": titles,
            "headers": headers,
            "narrative_texts": narrative_texts,
            "list_items": list_items,
            "page_numbers": page_numbers,
            "table_count": len(extracted_tables),
            "image_count": 0,
            "total_elements": len(elements),
        }
        
        print(f"✅ PaddleOCR + Unstructured 元素类解析完成，提取 {len(elements)} 个元素")
        print(f"   📊 统计: 0 张图片, {len(extracted_tables)} 个表格")
        print(f"   📝 结构: {len(titles)} 个标题, {len(narrative_texts)} 段正文")
        
        return {
            "text": full_text,
            "elements": elements,
            "structured_info": structured_info,
            "images": [],
            "tables": extracted_tables
        }
        
    except Exception as e:
        print(f"❌ PaddleOCR + Unstructured 元素类解析失败: {e}")
        print("🔄 回退到 PyMuPDF 模式")
        return process_pdf(pdf_path, output_dir)

